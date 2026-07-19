from __future__ import annotations

from pathlib import Path

import pytest

pytest.importorskip("docx")

import docx

from reclaim.ai.document_text import extract_text, is_supported_document


def test_is_supported_document_recognizes_known_extensions() -> None:
    for name in ["a.txt", "a.md", "a.docx", "a.pdf", "A.TXT"]:
        assert is_supported_document(Path(name)) is True
    assert is_supported_document(Path("a.jpg")) is False
    assert is_supported_document(Path("a.exe")) is False


def test_extract_text_reads_plain_text_files(tmp_path: Path) -> None:
    path = tmp_path / "notes.txt"
    path.write_text("hello world", encoding="utf-8")
    assert extract_text(path) == "hello world"


def test_extract_text_reads_markdown_files(tmp_path: Path) -> None:
    path = tmp_path / "notes.md"
    path.write_text("# Heading\n\nBody text.", encoding="utf-8")
    assert extract_text(path) == "# Heading\n\nBody text."


def test_extract_text_returns_none_for_unsupported_extension(tmp_path: Path) -> None:
    path = tmp_path / "image.jpg"
    path.write_bytes(b"not really a jpeg")
    assert extract_text(path) is None


def test_extract_text_returns_none_for_missing_file(tmp_path: Path) -> None:
    assert extract_text(tmp_path / "does_not_exist.txt") is None


def test_extract_text_returns_none_for_corrupt_docx(tmp_path: Path) -> None:
    path = tmp_path / "corrupt.docx"
    path.write_bytes(b"this is not a real docx zip archive")
    assert extract_text(path) is None


def test_extract_text_returns_none_for_corrupt_pdf(tmp_path: Path) -> None:
    path = tmp_path / "corrupt.pdf"
    path.write_bytes(b"%PDF-1.4 this is not a real pdf structure")
    assert extract_text(path) is None


def test_extract_text_real_docx_round_trips(tmp_path: Path) -> None:
    path = tmp_path / "real.docx"
    document = docx.Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    document.save(str(path))

    text = extract_text(path)
    assert text is not None
    assert "First paragraph." in text
    assert "Second paragraph." in text
